# -*- coding: utf-8 -*-
"""
Document Generation Service for HR Templates
Generates documents using organization templates with employee data
"""
import os
import json
from datetime import datetime, date
from pathlib import Path
from docx import Document
from docx.shared import Inches
import pandas as pd
from jinja2 import Template, Environment, FileSystemLoader
import streamlit as st
import re

# Try to import PDF libraries with proper error handling
PDF_AVAILABLE = False
WEASYPRINT_AVAILABLE = False
REPORTLAB_AVAILABLE = False

# Try WeasyPrint first
try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
    PDF_AVAILABLE = True
except (ImportError, OSError) as e:
    # WeasyPrint failed - this is common on Windows
    pass

# Try ReportLab as fallback
if not PDF_AVAILABLE:
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from io import BytesIO
        REPORTLAB_AVAILABLE = True
        PDF_AVAILABLE = True
    except ImportError:
        pass

class DocumentGenerator:
    """Document generator using organization templates"""
    
    def __init__(self):
        """Initialize document generator"""
        self.templates_dir = Path("templates")
        self.output_dir = Path("generated_documents")
        self.ensure_directories()
        
        # Initialize Jinja2 environment
        self.jinja_env = Environment(
            loader=FileSystemLoader(str(self.templates_dir)),
            autoescape=True
        )
        
        # Add custom filters
        self.jinja_env.filters['date_format'] = self.format_date
        self.jinja_env.filters['currency'] = self.format_currency
    
    def ensure_directories(self):
        """Create necessary directories"""
        self.templates_dir.mkdir(exist_ok=True)
        self.output_dir.mkdir(exist_ok=True)
        
        # Create subdirectories for different document types
        for doc_type in ['letters', 'contracts', 'certificates', 'reports']:
            (self.templates_dir / doc_type).mkdir(exist_ok=True)
            (self.output_dir / doc_type).mkdir(exist_ok=True)
    
    def get_available_templates(self):
        """Get list of available document templates"""
        templates = {}
        
        for category in ['letters', 'contracts', 'certificates', 'reports']:
            category_path = self.templates_dir / category
            if category_path.exists():
                templates[category] = []
                
                # Look for both .docx and .html templates
                for template_file in category_path.glob('*.docx'):
                    templates[category].append({
                        'name': template_file.stem,
                        'file': str(template_file),
                        'type': 'docx',
                        'category': category
                    })
                
                for template_file in category_path.glob('*.html'):
                    templates[category].append({
                        'name': template_file.stem,
                        'file': str(template_file),
                        'type': 'html',
                        'category': category
                    })
        
        return templates
    
    def format_date(self, date_value, format_string='%B %d, %Y'):
        """Format date for templates"""
        if isinstance(date_value, str):
            try:
                date_value = datetime.strptime(date_value, '%Y-%m-%d').date()
            except:
                return date_value
        
        if isinstance(date_value, (datetime, date)):
            return date_value.strftime(format_string)
        
        return str(date_value)
    
    def format_currency(self, amount, currency='USD'):
        """Format currency for templates"""
        try:
            amount = float(amount)
            if currency == 'USD':
                return f"${amount:,.2f}"
            else:
                return f"{amount:,.2f} {currency}"
        except:
            return str(amount)
    
    def get_sample_templates(self):
        """Create sample templates for common HR documents"""
        sample_templates = {
            'offer_letter': {
                'category': 'letters',
                'content': '''
<!DOCTYPE html>
<html>
<head>
    <title>Employment Offer Letter</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 40px; }
        .header { text-align: center; margin-bottom: 30px; }
        .content { line-height: 1.6; }
        .signature { margin-top: 40px; }
    </style>
</head>
<body>
    <div class="header">
        <h1>{{ company_name }}</h1>
        <h2>Employment Offer Letter</h2>
    </div>
    
    <div class="content">
        <p>{{ generation_date }}</p>
        
        <p>Dear {{ employee_name }},</p>
        
        <p>We are pleased to offer you the position of <strong>{{ position_title }}</strong> 
        in our {{ department }} department at {{ company_name }}.</p>
        
        <p><strong>Position Details:</strong></p>
        <ul>
            <li>Position: {{ position_title }}</li>
            <li>Department: {{ department }}</li>
            <li>Start Date: {{ start_date | date_format }}</li>
            <li>Salary: {{ salary | currency }}</li>
            <li>Employment Type: {{ employment_type }}</li>
        </ul>
        
        <p>This offer is contingent upon successful completion of background verification 
        and reference checks.</p>
        
        <p>Please confirm your acceptance by signing and returning this letter by 
        {{ response_deadline | date_format }}.</p>
        
        <p>We look forward to welcoming you to our team!</p>
        
        <div class="signature">
            <p>Sincerely,</p>
            <br><br>
            <p>{{ hr_manager_name }}<br>
            HR Manager<br>
            {{ company_name }}</p>
        </div>
        
        <hr style="margin-top: 40px;">
        <p>Employee Acceptance:</p>
        <p>I accept this offer of employment:</p>
        <br><br>
        <p>Signature: _________________ Date: _________</p>
    </div>
</body>
</html>
                '''
            },
            'termination_letter': {
                'category': 'letters',
                'content': '''
<!DOCTYPE html>
<html>
<head>
    <title>Employment Termination Letter</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 40px; }
        .header { text-align: center; margin-bottom: 30px; }
        .content { line-height: 1.6; }
    </style>
</head>
<body>
    <div class="header">
        <h1>{{ company_name }}</h1>
        <h2>Employment Termination Notice</h2>
    </div>
    
    <div class="content">
        <p>{{ generation_date }}</p>
        
        <p>Dear {{ employee_name }},</p>
        
        <p>This letter serves as formal notification that your employment with 
        {{ company_name }} will be terminated effective {{ termination_date | date_format }}.</p>
        
        <p><strong>Termination Details:</strong></p>
        <ul>
            <li>Employee ID: {{ employee_id }}</li>
            <li>Position: {{ position_title }}</li>
            <li>Department: {{ department }}</li>
            <li>Last Working Day: {{ last_working_day | date_format }}</li>
            <li>Reason: {{ termination_reason }}</li>
        </ul>
        
        {% if final_settlement %}
        <p><strong>Final Settlement:</strong></p>
        <ul>
            <li>Final Salary: {{ final_salary | currency }}</li>
            <li>Unused Leave: {{ unused_leave_days }} days ({{ unused_leave_amount | currency }})</li>
            <li>Total Amount: {{ total_settlement | currency }}</li>
        </ul>
        {% endif %}
        
        <p>Please return all company property including ID cards, equipment, and documents 
        before your last working day.</p>
        
        <p>Thank you for your service to {{ company_name }}.</p>
        
        <p>Sincerely,<br><br>
        {{ hr_manager_name }}<br>
        HR Manager</p>
    </div>
</body>
</html>
                '''
            },
            'experience_certificate': {
                'category': 'certificates',
                'content': '''
<!DOCTYPE html>
<html>
<head>
    <title>Experience Certificate</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 40px; }
        .header { text-align: center; margin-bottom: 40px; border: 2px solid #333; padding: 20px; }
        .content { line-height: 1.8; text-align: justify; }
        .signature { margin-top: 50px; float: right; }
    </style>
</head>
<body>
    <div class="header">
        <h1>{{ company_name }}</h1>
        <h2>EXPERIENCE CERTIFICATE</h2>
    </div>
    
    <div class="content">
        <p>TO WHOM IT MAY CONCERN</p>
        
        <p>This is to certify that <strong>{{ employee_name }}</strong> was employed with 
        {{ company_name }} from {{ start_date | date_format }} to {{ end_date | date_format }}.</p>
        
        <p>During the tenure of employment, {{ employee_name }} worked as 
        <strong>{{ position_title }}</strong> in the {{ department }} department.</p>
        
        <p>{{ employee_name }} has shown dedication, professionalism, and commitment to work. 
        {{ he_she | default('They') }} {{ was_were | default('were') }} a valuable team member and contributed 
        significantly to our organization.</p>
        
        <p>We wish {{ employee_name }} all the best for future endeavors.</p>
        
        <p>This certificate is issued upon request for employment purposes.</p>
        
        <div class="signature">
            <p>{{ generation_date }}</p>
            <br><br><br>
            <p>{{ hr_manager_name }}<br>
            HR Manager<br>
            {{ company_name }}</p>
            
            <p>Company Seal: ___________</p>
        </div>
    </div>
</body>
</html>
                '''
            }
        }
        
        return sample_templates
    
    def install_sample_templates(self):
        """Install sample templates to templates directory"""
        sample_templates = self.get_sample_templates()
        
        for template_name, template_info in sample_templates.items():
            category = template_info['category']
            content = template_info['content']
            
            # Create template file
            template_path = self.templates_dir / category / f"{template_name}.html"
            
            with open(template_path, 'w', encoding='utf-8') as f:
                f.write(content)
        
        return len(sample_templates)
    
    def generate_html_document(self, template_info, template_data, output_format='html'):
        """Generate HTML document from Jinja2 template"""
        try:
            # Store template data for PDF generation
            self._current_template_data = template_data
            
            # Load and render template
            template_path = Path(template_info['file'])
            template = self.jinja_env.get_template(f"{template_info['category']}/{template_path.name}")
            
            # Add system data
            template_data.update({
                'generation_date': datetime.now().strftime("%Y-%m-%d"),
                'generation_time': datetime.now().strftime("%H:%M:%S"),
                'company_name': template_data.get('company_name', 'Adanian Labs'),
            })
            
            # Render template
            rendered_content = template.render(**template_data)
            
            # Generate output filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            employee_name = template_data.get('employee_name', 'unknown').replace(' ', '_')
            
            if output_format == 'pdf':
                # Convert HTML to PDF
                return self.convert_html_to_pdf(rendered_content, template_info, employee_name, timestamp)
            else:
                # Save as HTML
                output_filename = f"{template_info['name']}_{employee_name}_{timestamp}.html"
                output_path = self.output_dir / template_info['category'] / output_filename
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(rendered_content)
                
                return {
                    'success': True,
                    'filename': output_filename,
                    'path': str(output_path),
                    'type': output_format,
                    'content': rendered_content
                }
        
        except Exception as e:
            # Add more detailed error information
            return {
                'success': False,
                'error': f"HTML document generation failed: {str(e)}",
                'template_data': template_data,
                'template_info': template_info
            }

    def convert_html_to_pdf(self, html_content, template_info, employee_name, timestamp):
        """Convert HTML content to PDF"""
        try:
            output_filename = f"{template_info['name']}_{employee_name}_{timestamp}.pdf"
            output_path = self.output_dir / template_info['category'] / output_filename
            
            if not PDF_AVAILABLE:
                return {
                    'success': False,
                    'error': "PDF generation requires additional libraries. Please install 'reportlab': pip install reportlab",
                    'template_data': {},
                    'template_info': template_info
                }
            
            # Try WeasyPrint first if available
            if WEASYPRINT_AVAILABLE:
                try:
                    # Add PDF-specific CSS for better printing
                    pdf_css = """
                        @page {
                            size: A4;
                            margin: 1in;
                        }
                        body {
                            font-family: Arial, sans-serif;
                            font-size: 12pt;
                            line-height: 1.4;
                            color: #333;
                        }
                        .header {
                            text-align: center;
                            margin-bottom: 30px;
                            border-bottom: 2px solid #333;
                            padding-bottom: 20px;
                        }
                        .signature {
                            margin-top: 40px;
                            page-break-inside: avoid;
                        }
                    """
                    
                    # Create HTML with embedded CSS
                    full_html = f"""
                    <!DOCTYPE html>
                    <html>
                    <head>
                        <meta charset="UTF-8">
                        <style>{pdf_css}</style>
                    </head>
                    <body>
                        {html_content.split('<body>')[1].split('</body>')[0] if '<body>' in html_content else html_content}
                    </body>
                    </html>
                    """
                    
                    # Generate PDF using weasyprint
                    pdf_bytes = HTML(string=full_html).write_pdf()
                    
                    # Save PDF file
                    with open(output_path, 'wb') as f:
                        f.write(pdf_bytes)
                    
                    return {
                        'success': True,
                        'filename': output_filename,
                        'path': str(output_path),
                        'type': 'pdf',
                        'content': pdf_bytes
                    }
                
                except Exception as weasy_error:
                    # Fall through to ReportLab
                    pass
            
            # Use ReportLab if WeasyPrint failed or isn't available
            if REPORTLAB_AVAILABLE:
                return self.create_reportlab_pdf(html_content, template_info, employee_name, timestamp)
            
            # If we get here, no PDF library worked
            return {
                'success': False,
                'error': "PDF generation failed. Please install ReportLab: pip install reportlab",
                'template_data': {},
                'template_info': template_info
            }
        
        except Exception as e:
            return {
                'success': False,
                'error': f"PDF generation failed: {str(e)}",
                'template_data': {},
                'template_info': template_info
            }

    def create_reportlab_pdf(self, html_content, template_info, employee_name, timestamp):
        """Create PDF using reportlab as fallback"""
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib.colors import black, blue
            import re
            
            output_filename = f"{template_info['name']}_{employee_name}_{timestamp}.pdf"
            output_path = self.output_dir / template_info['category'] / output_filename
            
            # Create PDF document
            doc = SimpleDocTemplate(str(output_path), pagesize=A4, 
                                  topMargin=inch, bottomMargin=inch, 
                                  leftMargin=inch, rightMargin=inch)
            
            # Get styles
            styles = getSampleStyleSheet()
            
            # Create custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=20,
                spaceAfter=30,
                alignment=1,  # Center
                textColor=blue
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=16,
                spaceAfter=20,
                alignment=1,  # Center
                textColor=black
            )
            
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontSize=12,
                spaceAfter=12,
                leftIndent=0,
                rightIndent=0
            )
            
            bold_style = ParagraphStyle(
                'CustomBold',
                parent=styles['Normal'],
                fontSize=12,
                spaceAfter=12,
                fontName='Helvetica-Bold'
            )
            
            # Instead of parsing HTML, create content directly from template data
            # Get the template data from the instance variable or extract from HTML
            template_data = getattr(self, '_current_template_data', {})
            
            # Build PDF content based on template type
            story = []
            
            if template_info['name'] == 'offer_letter':
                story = self.build_offer_letter_pdf(template_data, title_style, heading_style, normal_style, bold_style)
            elif template_info['name'] == 'termination_letter':
                story = self.build_termination_letter_pdf(template_data, title_style, heading_style, normal_style, bold_style)
            elif template_info['name'] == 'experience_certificate':
                story = self.build_experience_certificate_pdf(template_data, title_style, heading_style, normal_style, bold_style)
            else:
                # Fallback: extract from HTML
                content = self.extract_content_from_html(html_content, template_info['name'])
                story = self.build_generic_pdf(content, title_style, heading_style, normal_style, bold_style)
            
            # Build PDF
            doc.build(story)
            
            # Read the generated PDF
            with open(output_path, 'rb') as f:
                pdf_content = f.read()
            
            return {
                'success': True,
                'filename': output_filename,
                'path': str(output_path),
                'type': 'pdf',
                'content': pdf_content
            }
        
        except Exception as e:
            return {
                'success': False,
                'error': f"ReportLab PDF generation failed: {str(e)}",
                'template_data': {},
                'template_info': template_info
            }

    def build_offer_letter_pdf(self, data, title_style, heading_style, normal_style, bold_style):
        """Build offer letter PDF content from template data"""
        story = []
        
        # Company name as title
        company_name = data.get('company_name', 'Adanian Labs')
        story.append(Paragraph(company_name, title_style))
        story.append(Spacer(1, 12))
        
        # Document title
        story.append(Paragraph("Employment Offer Letter", heading_style))
        story.append(Spacer(1, 20))
        
        # Date
        story.append(Paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}", normal_style))
        story.append(Spacer(1, 12))
        
        # Greeting
        employee_name = data.get('employee_name', '[Employee Name]')
        story.append(Paragraph(f"Dear {employee_name},", normal_style))
        story.append(Spacer(1, 12))
        
        # Main content
        position_title = data.get('position_title', '[Position]')
        department = data.get('department', '[Department]')
        
        story.append(Paragraph(f"We are pleased to offer you the position of <b>{position_title}</b> in our {department} department at {company_name}.", normal_style))
        story.append(Spacer(1, 12))
        
        # Position details
        story.append(Paragraph("<b>Position Details:</b>", bold_style))
        story.append(Spacer(1, 8))
        
        details = [
            f"Position: {position_title}",
            f"Department: {department}",
            f"Start Date: {self.format_date(data.get('start_date', ''))}",
            f"Salary: {self.format_currency(data.get('salary', 0))}",
            f"Employment Type: {data.get('employment_type', '[Type]')}"
        ]
        
        for detail in details:
            story.append(Paragraph(f"• {detail}", normal_style))
        
        story.append(Spacer(1, 12))
        
        # Additional content
        story.append(Paragraph("This offer is contingent upon successful completion of background verification and reference checks.", normal_style))
        story.append(Spacer(1, 8))
        
        response_deadline = self.format_date(data.get('response_deadline', ''))
        story.append(Paragraph(f"Please confirm your acceptance by signing and returning this letter by {response_deadline}.", normal_style))
        story.append(Spacer(1, 8))
        
        story.append(Paragraph("We look forward to welcoming you to our team!", normal_style))
        story.append(Spacer(1, 20))
        
        # Signature section
        story.append(Paragraph("Sincerely,", normal_style))
        story.append(Spacer(1, 24))
        
        hr_manager = data.get('hr_manager_name', '[HR Manager]')
        story.append(Paragraph(hr_manager, normal_style))
        story.append(Paragraph("HR Manager", normal_style))
        story.append(Paragraph(company_name, normal_style))
        
        # Acceptance section
        story.append(Spacer(1, 30))
        story.append(Paragraph("_" * 60, normal_style))
        story.append(Spacer(1, 12))
        story.append(Paragraph("<b>Employee Acceptance:</b>", bold_style))
        story.append(Paragraph("I accept this offer of employment:", normal_style))
        story.append(Spacer(1, 20))
        story.append(Paragraph("Signature: _________________________ Date: __________", normal_style))
        
        return story

    def build_termination_letter_pdf(self, data, title_style, heading_style, normal_style, bold_style):
        """Build termination letter PDF content from template data"""
        story = []
        
        # Company name as title
        company_name = data.get('company_name', 'Adanian Labs')
        story.append(Paragraph(company_name, title_style))
        story.append(Spacer(1, 12))
        
        # Document title
        story.append(Paragraph("Employment Termination Notice", heading_style))
        story.append(Spacer(1, 20))
        
        # Date
        story.append(Paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}", normal_style))
        story.append(Spacer(1, 12))
        
        # Greeting
        employee_name = data.get('employee_name', '[Employee Name]')
        story.append(Paragraph(f"Dear {employee_name},", normal_style))
        story.append(Spacer(1, 12))
        
        # Main content
        termination_date = self.format_date(data.get('termination_date', ''))
        story.append(Paragraph(f"This letter serves as formal notification that your employment with {company_name} will be terminated effective {termination_date}.", normal_style))
        story.append(Spacer(1, 12))
        
        # Termination details
        story.append(Paragraph("<b>Termination Details:</b>", bold_style))
        story.append(Spacer(1, 8))
        
        details = [
            f"Employee ID: {data.get('employee_id', '[ID]')}",
            f"Position: {data.get('position_title', '[Position]')}",
            f"Department: {data.get('department', '[Department]')}",
            f"Last Working Day: {self.format_date(data.get('last_working_day', ''))}",
            f"Reason: {data.get('termination_reason', '[Reason]')}"
        ]
        
        for detail in details:
            story.append(Paragraph(f"• {detail}", normal_style))
        
        story.append(Spacer(1, 12))
        
        # Final settlement if provided
        if data.get('final_settlement'):
            story.append(Paragraph("<b>Final Settlement:</b>", bold_style))
            story.append(Spacer(1, 8))
            
            settlement_details = [
                f"Final Salary: {self.format_currency(data.get('final_salary', 0))}",
                f"Unused Leave: {data.get('unused_leave_days', 0)} days ({self.format_currency(data.get('unused_leave_amount', 0))})",
                f"Total Amount: {self.format_currency(data.get('total_settlement', 0))}"
            ]
            
            for detail in settlement_details:
                story.append(Paragraph(f"• {detail}", normal_style))
            
            story.append(Spacer(1, 12))
        
        # Additional content
        story.append(Paragraph("Please return all company property including ID cards, equipment, and documents before your last working day.", normal_style))
        story.append(Spacer(1, 8))
        
        story.append(Paragraph(f"Thank you for your service to {company_name}.", normal_style))
        story.append(Spacer(1, 20))
        
        # Signature section
        story.append(Paragraph("Sincerely,", normal_style))
        story.append(Spacer(1, 24))
        
        hr_manager = data.get('hr_manager_name', '[HR Manager]')
        story.append(Paragraph(hr_manager, normal_style))
        story.append(Paragraph("HR Manager", normal_style))
        
        return story

    def build_experience_certificate_pdf(self, data, title_style, heading_style, normal_style, bold_style):
        """Build experience certificate PDF content from template data"""
        story = []
        
        # Company name as title
        company_name = data.get('company_name', 'Adanian Labs')
        story.append(Paragraph(company_name, title_style))
        story.append(Spacer(1, 12))
        
        # Document title
        story.append(Paragraph("EXPERIENCE CERTIFICATE", heading_style))
        story.append(Spacer(1, 30))
        
        # TO WHOM IT MAY CONCERN
        concern_para = Paragraph("TO WHOM IT MAY CONCERN", bold_style)
        concern_para.alignment = 1  # Center
        story.append(concern_para)
        story.append(Spacer(1, 20))
        
        # Main content
        employee_name = data.get('employee_name', '[Employee Name]')
        start_date = self.format_date(data.get('start_date', ''))
        end_date = self.format_date(data.get('end_date', ''))
        
        story.append(Paragraph(f"This is to certify that <b>{employee_name}</b> was employed with {company_name} from {start_date} to {end_date}.", normal_style))
        story.append(Spacer(1, 12))
        
        position_title = data.get('position_title', '[Position]')
        department = data.get('department', '[Department]')
        
        story.append(Paragraph(f"During the tenure of employment, {employee_name} worked as <b>{position_title}</b> in the {department} department.", normal_style))
        story.append(Spacer(1, 12))
        
        he_she = data.get('he_she', 'They')
        was_were = data.get('was_were', 'were')
        
        story.append(Paragraph(f"{employee_name} has shown dedication, professionalism, and commitment to work. {he_she} {was_were} a valuable team member and contributed significantly to our organization.", normal_style))
        story.append(Spacer(1, 12))
        
        story.append(Paragraph(f"We wish {employee_name} all the best for future endeavors.", normal_style))
        story.append(Spacer(1, 12))
        
        story.append(Paragraph("This certificate is issued upon request for employment purposes.", normal_style))
        story.append(Spacer(1, 40))
        
        # Signature section (right aligned)
        date_para = Paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}", normal_style)
        date_para.alignment = 2  # Right align
        story.append(date_para)
        story.append(Spacer(1, 30))
        
        hr_manager = data.get('hr_manager_name', '[HR Manager]')
        sig_para = Paragraph(hr_manager, normal_style)
        sig_para.alignment = 2
        story.append(sig_para)
        
        mgr_para = Paragraph("HR Manager", normal_style)
        mgr_para.alignment = 2
        story.append(mgr_para)
        
        comp_para = Paragraph(company_name, normal_style)
        comp_para.alignment = 2
        story.append(comp_para)
        
        story.append(Spacer(1, 12))
        seal_para = Paragraph("Company Seal: ___________", normal_style)
        seal_para.alignment = 2
        story.append(seal_para)
        
        return story

    def build_generic_pdf(self, content, title_style, heading_style, normal_style, bold_style):
        """Build generic PDF content from extracted HTML content"""
        story = []
        
        # Add company name as title
        story.append(Paragraph(content['company_name'], title_style))
        story.append(Spacer(1, 12))
        
        # Add document title
        story.append(Paragraph(content['document_title'], heading_style))
        story.append(Spacer(1, 20))
        
        # Add date
        if content['date']:
            story.append(Paragraph(f"Date: {content['date']}", normal_style))
            story.append(Spacer(1, 12))
        
        # Add greeting
        if content['greeting']:
            story.append(Paragraph(content['greeting'], normal_style))
            story.append(Spacer(1, 12))
        
        # Add main content paragraphs
        for paragraph in content['main_content']:
            if paragraph.strip():
                story.append(Paragraph(paragraph.strip(), normal_style))
                story.append(Spacer(1, 8))
        
        # Add position details section
        if content['details']:
            story.append(Spacer(1, 12))
            story.append(Paragraph("<b>Details:</b>", bold_style))
            for detail in content['details']:
                story.append(Paragraph(f"• {detail}", normal_style))
            story.append(Spacer(1, 12))
        
        # Add signature section
        story.append(Spacer(1, 20))
        story.append(Paragraph("Sincerely,", normal_style))
        story.append(Spacer(1, 24))
        story.append(Paragraph(content['hr_manager'], normal_style))
        story.append(Paragraph("HR Manager", normal_style))
        story.append(Paragraph(content['company_name'], normal_style))
        
        return story

    def extract_content_from_html(self, html_content, template_name):
        """Extract structured content from HTML for PDF generation"""
        import re
        
        # Remove HTML tags and extract text
        text_content = re.sub('<[^<]+?>', '\n', html_content)
        text_content = re.sub(r'\s+', ' ', text_content)
        text_content = text_content.replace('&nbsp;', ' ').replace('&amp;', '&')
        
        # Initialize content structure
        content = {
            'company_name': 'Adanian Labs',
            'document_title': 'HR Document',
            'date': '',
            'greeting': '',
            'main_content': [],
            'details': [],
            'additional_content': [],
            'hr_manager': '[HR Manager]'
        }
        
        # Split into lines and process
        lines = [line.strip() for line in text_content.split('\n') if line.strip()]
        
        # Template-specific parsing
        if template_name == 'offer_letter':
            content.update(self.parse_offer_letter_content(lines))
        elif template_name == 'termination_letter':
            content.update(self.parse_termination_letter_content(lines))
        elif template_name == 'experience_certificate':
            content.update(self.parse_experience_certificate_content(lines))
        else:
            content['main_content'] = lines
        
        return content

    def parse_offer_letter_content(self, lines):
        """Parse offer letter specific content"""
        content = {
            'document_title': 'Employment Offer Letter',
            'main_content': [],
            'details': [],
            'additional_content': []
        }
        
        detail_section = False
        signature_section = False
        
        for line in lines:
            # Skip CSS and styling text
            if any(skip in line.lower() for skip in ['font-family', 'margin:', 'text-align', 'line-height', '.header', '.content', '.signature']):
                continue
                
            # Extract company name (usually first meaningful line)
            if 'adanian labs' in line.lower() or 'company' in line.lower():
                if not content.get('company_name'):
                    content['company_name'] = line
                continue
            
            # Skip repeated document title
            if 'employment offer letter' in line.lower():
                continue
                
            # Extract date (YYYY-MM-DD format)
            if re.match(r'\d{4}-\d{2}-\d{2}', line):
                content['date'] = line
                continue
                
            # Extract greeting
            if line.lower().startswith('dear '):
                content['greeting'] = line
                continue
                
            # Position details section
            if 'position details:' in line.lower():
                detail_section = True
                continue
            elif 'position:' in line or 'department:' in line or 'start date:' in line or 'salary:' in line or 'employment type:' in line:
                if detail_section:
                    content['details'].append(line)
                    continue
                    
            # End of details section
            if detail_section and ('this offer is contingent' in line.lower() or 'please confirm' in line.lower()):
                detail_section = False
                
            # Signature section
            if 'sincerely' in line.lower():
                signature_section = True
                continue
            elif signature_section and ('hr manager' not in line.lower() and 'employee acceptance' not in line.lower() and 'signature:' not in line.lower()):
                content['hr_manager'] = line
                signature_section = False
                continue
                
            # Skip signature and acceptance lines
            if any(skip in line.lower() for skip in ['hr manager', 'employee acceptance', 'i accept this offer', 'signature:', '____']):
                continue
                
            # Add to main content
            if not detail_section and not signature_section and line:
                content['main_content'].append(line)
        
        return content

    def parse_termination_letter_content(self, lines):
        """Parse termination letter specific content"""
        content = {
            'document_title': 'Employment Termination Notice',
            'main_content': [],
            'details': [],
            'additional_content': []
        }
        
        detail_section = False
        settlement_section = False
        
        for line in lines:
            # Skip CSS
            if any(skip in line.lower() for skip in ['font-family', 'margin:', 'text-align', 'line-height']):
                continue
                
            # Extract company name
            if 'adanian labs' in line.lower():
                content['company_name'] = line
                continue
                
            # Skip repeated title
            if 'employment termination' in line.lower():
                continue
                
            # Extract date
            if re.match(r'\d{4}-\d{2}-\d{2}', line):
                content['date'] = line
                continue
                
            # Extract greeting
            if line.lower().startswith('dear '):
                content['greeting'] = line
                continue
                
            # Termination details section
            if 'termination details:' in line.lower():
                detail_section = True
                continue
            elif detail_section and ('employee id:' in line.lower() or 'position:' in line.lower() or 'department:' in line.lower() or 'last working day:' in line.lower() or 'reason:' in line.lower()):
                content['details'].append(line)
                continue
                
            # Final settlement section
            if 'final settlement:' in line.lower():
                settlement_section = True
                detail_section = False
                continue
            elif settlement_section and ('final salary:' in line.lower() or 'unused leave:' in line.lower() or 'total amount:' in line.lower()):
                content['details'].append(line)
                continue
                
            # End of sections
            if 'please return all company' in line.lower():
                detail_section = False
                settlement_section = False
                
            # HR Manager
            if not any(skip in line.lower() for skip in ['hr manager', 'sincerely']) and line and not detail_section and not settlement_section:
                if 'thank you for your service' not in line.lower():
                    content['main_content'].append(line)
        
        return content

    def parse_experience_certificate_content(self, lines):
        """Parse experience certificate specific content"""
        content = {
            'document_title': 'EXPERIENCE CERTIFICATE',
            'main_content': [],
            'details': [],
            'additional_content': []
        }
        
        for line in lines:
            # Skip CSS
            if any(skip in line.lower() for skip in ['font-family', 'margin:', 'text-align', 'line-height']):
                continue
                
            # Extract company name
            if 'adanian labs' in line.lower():
                content['company_name'] = line
                continue
                
            # Skip repeated title
            if 'experience certificate' in line.lower():
                continue
                
            # Extract date
            if re.match(r'\d{4}-\d{2}-\d{2}', line):
                content['date'] = line
                continue
                
            # Skip signature elements
            if any(skip in line.lower() for skip in ['hr manager', 'company seal']):
                continue
                
            # Add to main content
            if line:
                content['main_content'].append(line)
        
        return content

    def generate_docx_document(self, template_info, template_data):
        """Generate Word document from template"""
        try:
            # Load template document
            doc = Document(template_info['file'])
            
            # Replace placeholders in paragraphs
            for paragraph in doc.paragraphs:
                for key, value in template_data.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, str(value))
            
            # Replace placeholders in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in template_data.items():
                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in cell.text:
                                cell.text = cell.text.replace(placeholder, str(value))
            
            # Generate output filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            employee_name = template_data.get('employee_name', 'unknown').replace(' ', '_')
            output_filename = f"{template_info['name']}_{employee_name}_{timestamp}.docx"
            output_path = self.output_dir / template_info['category'] / output_filename
            
            # Save document
            doc.save(str(output_path))
            
            # Read the saved document for download
            with open(output_path, 'rb') as f:
                docx_content = f.read()
            
            return {
                'success': True,
                'filename': output_filename,
                'path': str(output_path),
                'type': 'docx',
                'content': docx_content
            }
        
        except Exception as e:
            return {
                'success': False,
                'error': f"DOCX document generation failed: {str(e)}",
                'template_data': template_data,
                'template_info': template_info
            }

    def create_docx_from_html_template(self, template_info, template_data):
        """Create a DOCX document from HTML template data"""
        try:
            # Create a new Word document
            doc = Document()
            
            # Add company header
            company_name = template_data.get('company_name', 'Adanian Labs')
            header = doc.add_heading(company_name, 0)
            header.alignment = 1  # Center alignment
            
            # Add document title based on template type
            if template_info['name'] == 'offer_letter':
                title = doc.add_heading('Employment Offer Letter', level=1)
                title.alignment = 1
                self.add_offer_letter_content(doc, template_data)
            elif template_info['name'] == 'termination_letter':
                title = doc.add_heading('Employment Termination Notice', level=1)
                title.alignment = 1
                self.add_termination_letter_content(doc, template_data)
            elif template_info['name'] == 'experience_certificate':
                title = doc.add_heading('EXPERIENCE CERTIFICATE', level=1)
                title.alignment = 1
                self.add_experience_certificate_content(doc, template_data)
            else:
                title = doc.add_heading('HR Document', level=1)
                title.alignment = 1
                self.add_generic_content(doc, template_data)
            
            # Generate output filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            employee_name = template_data.get('employee_name', 'unknown').replace(' ', '_')
            output_filename = f"{template_info['name']}_{employee_name}_{timestamp}.docx"
            output_path = self.output_dir / template_info['category'] / output_filename
            
            # Save document
            doc.save(str(output_path))
            
            # Read the saved document for download
            with open(output_path, 'rb') as f:
                docx_content = f.read()
            
            return {
                'success': True,
                'filename': output_filename,
                'path': str(output_path),
                'type': 'docx',
                'content': docx_content
            }
        
        except Exception as e:
            return {
                'success': False,
                'error': f"DOCX creation from HTML template failed: {str(e)}",
                'template_data': template_data,
                'template_info': template_info
            }

    def add_offer_letter_content(self, doc, data):
        """Add offer letter content to DOCX"""
        # Date
        doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph("")
        
        # Greeting
        doc.add_paragraph(f"Dear {data.get('employee_name', '[Employee Name]')},")
        doc.add_paragraph("")
        
        # Main content
        doc.add_paragraph(f"We are pleased to offer you the position of {data.get('position_title', '[Position]')} in our {data.get('department', '[Department]')} department at {data.get('company_name', 'Adanian Labs')}.")
        doc.add_paragraph("")
        
        # Position details
        doc.add_heading('Position Details:', level=2)
        details = doc.add_paragraph()
        details.add_run(f"• Position: {data.get('position_title', '[Position]')}\n")
        details.add_run(f"• Department: {data.get('department', '[Department]')}\n")
        details.add_run(f"• Start Date: {self.format_date(data.get('start_date', ''))}\n")
        details.add_run(f"• Salary: {self.format_currency(data.get('salary', 0))}\n")
        details.add_run(f"• Employment Type: {data.get('employment_type', '[Type]')}")
        
        doc.add_paragraph("")
        doc.add_paragraph("This offer is contingent upon successful completion of background verification and reference checks.")
        doc.add_paragraph("")
        doc.add_paragraph(f"Please confirm your acceptance by signing and returning this letter by {self.format_date(data.get('response_deadline', ''))}.")
        doc.add_paragraph("")
        doc.add_paragraph("We look forward to welcoming you to our team!")
        doc.add_paragraph("")
        
        # Signature
        doc.add_paragraph("Sincerely,")
        doc.add_paragraph("")
        doc.add_paragraph("")
        doc.add_paragraph(f"{data.get('hr_manager_name', '[HR Manager]')}")
        doc.add_paragraph("HR Manager")
        doc.add_paragraph(f"{data.get('company_name', 'Adanian Labs')}")
        
        # Acceptance section
        doc.add_paragraph("_" * 50)
        doc.add_paragraph("")
        doc.add_paragraph("Employee Acceptance:")
        doc.add_paragraph("I accept this offer of employment:")
        doc.add_paragraph("")
        doc.add_paragraph("Signature: _________________ Date: _________")

    def add_termination_letter_content(self, doc, data):
        """Add termination letter content to DOCX"""
        # Date
        doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph("")
        
        # Greeting
        doc.add_paragraph(f"Dear {data.get('employee_name', '[Employee Name]')},")
        doc.add_paragraph("")
        
        # Main content
        doc.add_paragraph(f"This letter serves as formal notification that your employment with {data.get('company_name', 'Adanian Labs')} will be terminated effective {self.format_date(data.get('termination_date', ''))}.")
        doc.add_paragraph("")
        
        # Termination details
        doc.add_heading('Termination Details:', level=2)
        details = doc.add_paragraph()
        details.add_run(f"• Employee ID: {data.get('employee_id', '[ID]')}\n")
        details.add_run(f"• Position: {data.get('position_title', '[Position]')}\n")
        details.add_run(f"• Department: {data.get('department', '[Department]')}\n")
        details.add_run(f"• Last Working Day: {self.format_date(data.get('last_working_day', ''))}\n")
        details.add_run(f"• Reason: {data.get('termination_reason', '[Reason]')}")
        
        # Final settlement if provided
        if data.get('final_settlement'):
            doc.add_paragraph("")
            doc.add_heading('Final Settlement:', level=2)
            settlement = doc.add_paragraph()
            settlement.add_run(f"• Final Salary: {self.format_currency(data.get('final_salary', 0))}\n")
            settlement.add_run(f"• Unused Leave: {data.get('unused_leave_days', 0)} days ({self.format_currency(data.get('unused_leave_amount', 0))})\n")
            settlement.add_run(f"• Total Amount: {self.format_currency(data.get('total_settlement', 0))}")
        
        doc.add_paragraph("")
        doc.add_paragraph("Please return all company property including ID cards, equipment, and documents before your last working day.")
        doc.add_paragraph("")
        doc.add_paragraph(f"Thank you for your service to {data.get('company_name', 'Adanian Labs')}.")
        doc.add_paragraph("")
        
        # Signature
        doc.add_paragraph("Sincerely,")
        doc.add_paragraph("")
        doc.add_paragraph("")
        doc.add_paragraph(f"{data.get('hr_manager_name', '[HR Manager]')}")
        doc.add_paragraph("HR Manager")

    def add_experience_certificate_content(self, doc, data):
        """Add experience certificate content to DOCX"""
        # TO WHOM IT MAY CONCERN
        para = doc.add_paragraph("TO WHOM IT MAY CONCERN")
        para.alignment = 1  # Center alignment
        doc.add_paragraph("")
        
        # Main content
        doc.add_paragraph(f"This is to certify that {data.get('employee_name', '[Employee Name]')} was employed with {data.get('company_name', 'Adanian Labs')} from {self.format_date(data.get('start_date', ''))} to {self.format_date(data.get('end_date', ''))}.")
        doc.add_paragraph("")
        
        doc.add_paragraph(f"During the tenure of employment, {data.get('employee_name', '[Employee Name]')} worked as {data.get('position_title', '[Position]')} in the {data.get('department', '[Department]')} department.")
        doc.add_paragraph("")
        
        doc.add_paragraph(f"{data.get('employee_name', '[Employee Name]')} has shown dedication, professionalism, and commitment to work. {data.get('he_she', 'They')} {data.get('was_were', 'were')} a valuable team member and contributed significantly to our organization.")
        doc.add_paragraph("")
        
        doc.add_paragraph(f"We wish {data.get('employee_name', '[Employee Name]')} all the best for future endeavors.")
        doc.add_paragraph("")
        
        doc.add_paragraph("This certificate is issued upon request for employment purposes.")
        doc.add_paragraph("")
        doc.add_paragraph("")
        
        # Signature section (right aligned)
        signature_para = doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        signature_para.alignment = 2  # Right alignment
        doc.add_paragraph("")
        doc.add_paragraph("")
        
        sig_para = doc.add_paragraph(f"{data.get('hr_manager_name', '[HR Manager]')}")
        sig_para.alignment = 2
        sig_para = doc.add_paragraph("HR Manager")
        sig_para.alignment = 2
        sig_para = doc.add_paragraph(f"{data.get('company_name', 'Adanian Labs')}")
        sig_para.alignment = 2
        doc.add_paragraph("")
        sig_para = doc.add_paragraph("Company Seal: ___________")
        sig_para.alignment = 2

    def add_generic_content(self, doc, data):
        """Add generic content to DOCX"""
        doc.add_paragraph(f"Employee: {data.get('employee_name', '[Employee Name]')}")
        doc.add_paragraph(f"Position: {data.get('position_title', '[Position]')}")
        doc.add_paragraph(f"Department: {data.get('department', '[Department]')}")
        doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")

    def generate_document(self, template_name, template_data, output_format='docx'):
        """Generate document from template"""
        try:
            # Find template
            templates = self.get_available_templates()
            template_info = None
            
            for category, template_list in templates.items():
                for template in template_list:
                    if template['name'] == template_name:
                        template_info = template
                        break
                if template_info:
                    break
            
            if not template_info:
                raise Exception(f"Template '{template_name}' not found")
            
            # Check PDF availability and provide user-friendly error
            if output_format == 'pdf' and not PDF_AVAILABLE:
                return {
                    'success': False,
                    'error': """PDF generation is not available. To enable PDF generation, please install one of these libraries:

**Option 1 (Recommended for most cases):**
```
pip install reportlab
```

**Option 2 (Better formatting but requires system dependencies):**
```
pip install weasyprint
```

Note: WeasyPrint may require additional system libraries on Windows. ReportLab is easier to install and works well for most use cases.

After installation, restart your Streamlit application.""",
                    'template_data': template_data,
                    'template_info': template_info
                }
            
            # Generate based on requested output format
            if output_format == 'docx':
                # If we have a docx template, use it; otherwise create from HTML template
                if template_info['type'] == 'docx':
                    return self.generate_docx_document(template_info, template_data)
                else:
                    # Create DOCX from HTML template data
                    return self.create_docx_from_html_template(template_info, template_data)
            else:
                # Generate HTML or PDF
                return self.generate_html_document(template_info, template_data, output_format)
        
        except Exception as e:
            raise Exception(f"Document generation failed: {str(e)}")
